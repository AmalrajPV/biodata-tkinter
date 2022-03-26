from tkinter import *
from tkinter.filedialog import askopenfilename
from tkinter import messagebox
from docx import Document
from docx.shared import Inches


def open_img():
    path = askopenfilename(
        filetypes=[('Image Files', ['*.jpeg', '*.png', '*jpg'])])
    if path:
        img_path.set(path)


def is_valid():
    if name.get() != '' and email.get() != '' and dob.get() != '' and address.get() != '' and \
            qualification.get() != '':
        return True
    else:
        messagebox.showerror('Field Empty', 'All fields are mandatory')


def create_doc():
    if is_valid():
        document = Document()
        document.add_heading('BIODATA', 0)
        info_table = document.add_table(rows=0, cols=3)
        data = (
            ('Name', name.get()),
            ('Email', email.get()),
            ('Date of Birth', dob.get()),
            ('Gender', gender.get()),
            ('Address', address.get()),
            ('Qualification', qualification.get())
        )
        for n, d in data:
            row = info_table.add_row().cells
            row[0].text = n
            row[1].text = d
            for j in row[0].paragraphs:
                for run in j.runs:
                    run.font.bold = True

        c1 = info_table.cell(0, 2)
        c2 = info_table.cell(1, 2)
        c1.merge(c2)
        c1 = info_table.cell(1, 2)
        c2 = info_table.cell(2, 2)
        c1.merge(c2)

        c1 = info_table.cell(3, 1)
        c2 = info_table.cell(3, 2)
        c1.merge(c2)
        c1 = info_table.cell(4, 1)
        c2 = info_table.cell(4, 2)
        c1.merge(c2)
        c1 = info_table.cell(5, 1)
        c2 = info_table.cell(5, 2)
        c1.merge(c2)

        cell = info_table.rows[0].cells
        para = cell[2].paragraphs[0]
        run = para.add_run()
        run.add_picture(img_path.get(), width=Inches(2), height=Inches(2))
        # document.add_picture(img_path.get(), width=Inches(2), height=Inches(2))
        for r in info_table.rows:
            r.height = Inches(.7)
        document.save(f'{name.get()}.docx')
        messagebox.showinfo('Success', 'Your biodata created successfuly')


root = Tk()
root.geometry('500x600')
root.title('Biodata')
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=1)
img_path = StringVar()
name = StringVar()
email = StringVar()
dob = StringVar()
gender = StringVar()
address = StringVar()
qualification = StringVar()
gender.set('male')

f = ('Times', 14)
f2 = ('Times', 22, 'bold')

# title
title_l = Label(root, text="BIODATA", font=f2)
title_l.grid(column=0, row=0, columnspan=2, pady=10)

# Name
name_l = Label(root, text="Name", font=f)
name_l.grid(column=0, row=1, pady=10, sticky=W, ipadx=30)
name_e = Entry(root, font=f, textvariable=name)
name_e.grid(column=1, row=1, pady=10, sticky=EW, padx=30)

# Email
email_l = Label(root, text="Email", font=f)
email_l.grid(column=0, row=2, pady=10, sticky=W, ipadx=30)
email_e = Entry(root, font=f, textvariable=email)
email_e.grid(column=1, row=2, pady=10, sticky=EW, padx=30)

# Dob
dob_l = Label(root, text="Date of Birth", font=f)
dob_l.grid(column=0, row=3, pady=10, sticky=W, ipadx=30)
dob_e = Entry(root, font=f, textvariable=dob)
dob_e.grid(column=1, row=3, pady=10, sticky=EW, padx=30)

# Gender
gender_l = Label(root, text="Gender", font=f)
gender_l.grid(column=0, row=4, pady=10, sticky=W, ipadx=30)
gender_frame = LabelFrame(root, padx=10, pady=10)
gender_frame.grid(column=1, row=4)
mr = Radiobutton(gender_frame, text='Male', value='male', variable=gender, font=('Times', 10))
mr.grid(row=0, column=0, padx=5)
fr = Radiobutton(gender_frame, text='Female', value='female', variable=gender, font=('Times', 10))
fr.grid(row=0, column=1, padx=5)
o_r = Radiobutton(gender_frame, text='Other', value='other', variable=gender, font=('Times', 10))
o_r.grid(row=0, column=2, padx=5)

# Address
address_l = Label(root, text="Address", font=f)
address_l.grid(column=0, row=5, pady=10, sticky=W, ipadx=30)
address_e = Entry(root, font=f, textvariable=address)
address_e.grid(column=1, row=5, pady=10, sticky=EW, padx=30)

# Qualification
qualification_l = Label(root, text="Qualification", font=f)
qualification_l.grid(column=0, row=6, pady=10, sticky=W, ipadx=30)
qualificatione_e = Entry(root, font=f, textvariable=qualification)
qualificatione_e.grid(column=1, row=6, pady=10, sticky=EW, padx=30)

# Images
image_l = Label(root, text="Image", font=f)
image_l.grid(column=0, row=7, pady=10, sticky=W, ipadx=30)
img_f = LabelFrame(root)
img_f.grid(column=1, row=7, sticky=EW, padx=30)
img_f.grid_columnconfigure(0, weight=3)
img_f.grid_columnconfigure(1, weight=1)
path_e = Entry(img_f, state="disabled", textvariable=img_path, relief='flat')
path_e.grid(column=0, row=0, ipady=5, sticky=EW)
add_img = Button(img_f, text="Browse", command=open_img, font=('Times', 11), relief='solid')
add_img.grid(column=1, row=0, sticky=EW)

submit = Button(root, text="Create", font=f, command=create_doc)
submit.grid(row=8, column=1, padx=30, pady=10, sticky=EW)

root.mainloop()
